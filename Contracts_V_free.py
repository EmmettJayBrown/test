from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document(r'C:\Users\cavin396\OneDrive - BLUEHOLE\바탕 화면\업무\02. FU 업무\00. 계약서 양식\01. 성우계약서_개인\크래프톤\성우계약서_크래프톤_개인.docx')


customer_name = "전재현"
gamename = "TERA"
date_issue = "2022년 05월 22일"
workname = "마일스톤 리뷰 영상"
date_start = "2022년 05월 22일"
date_end = "2022년 07월 22일"
money_num = "1,000,000"
money_kor = "일백만"
customer_birth = "1993년 02월 15일"
customer_hp = "010-8639-3431"
customer_address1 = "경기도 군포시 산본로432번길 25"
customer_address2 = "1214동 201호"
date_inspect = "2022년 06월 22일"



for p in doc.paragraphs:
    if "(계약자 이름)" in p.text:
         p.text = p.text.replace("(계약자 이름)",customer_name)
    if "(게임명)" in p.text:
         p.text = p.text.replace("(게임명)",gamename)
    if "(용역일)" in p.text:
         p.text = p.text.replace("(용역일)",date_issue)
    if "(담당업무)" in p.text:
         p.text = p.text.replace("(담당업무)",workname)
    if "(계약 시작일)" in p.text:
         p.text = p.text.replace("(계약 시작일)",date_start)
    if "(계약 종료일)" in p.text:
         p.text = p.text.replace("(계약 종료일)",date_end)
    if "(계약금/숫자)" in p.text:
         p.text = p.text.replace("(계약금/숫자)",money_num)
    if "(계약금/한글)" in p.text:
         p.text = p.text.replace("(계약금/한글)",money_kor)
    if "(계약자 생년월일)" in p.text:
         p.text = p.text.replace("(계약자 생년월일)",customer_birth)
    if "(계약자 전화번호)" in p.text:
         p.text = p.text.replace("(계약자 전화번호)",customer_hp)
    if "(계약자 도로명 주소)" in p.text:
         p.text = p.text.replace("(계약자 도로명 주소)",customer_address1)
    if "(계약자 상세 주소)" in p.text:
         p.text = p.text.replace("(계약자 상세 주소)",customer_address2)
    if "(검수일)" in p.text:
         p.text = p.text.replace("(검수일)",date_inspect)


table = doc.tables[0]

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if '(계약자 이름)' in paragraph.text:
                    paragraph.text = paragraph.text.replace("(계약자 이름)",customer_name)
                if '(계약자 생년월일)' in paragraph.text:
                    paragraph.text = paragraph.text.replace("(계약자 생년월일)",customer_birth)
                if '(계약자 전화번호)' in paragraph.text:
                    paragraph.text = paragraph.text.replace("(계약자 전화번호)",customer_hp)
                if '(계약자 도로명 주소)' in paragraph.text:
                    paragraph.text = paragraph.text.replace("(계약자 도로명 주소)",customer_address1)
                if '(계약자 상세 주소)' in paragraph.text:
                    paragraph.text = paragraph.text.replace("(계약자 상세 주소)",customer_address2)


table = doc.tables[1]

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if '(게임명)' in paragraph.text:
                    paragraph.text = paragraph.text.replace("(게임명)",gamename)
                if '(담당업무)' in paragraph.text:
                    paragraph.text = paragraph.text.replace("(담당업무)",workname)




doc.save(r'C:\Users\cavin396\OneDrive - BLUEHOLE\바탕 화면\업무\02. FU 업무\test\성우계약서_크래프톤_개인_test_49.docx')

